"""
Servicio de generación de documentos PDF.
Rellena plantillas .docx con datos reales, convierte a PDF con LibreOffice
y superpone el sello cuando es requerido.
"""
import io
import os
import subprocess
import tempfile
from datetime import datetime

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from pypdf import PdfReader, PdfWriter

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "..", "templates")


def _ruta(nombre: str) -> str:
    return os.path.join(TEMPLATES_DIR, nombre)


def _reemplazar_variables(doc: Document, variables: dict) -> Document:
    """
    Reemplaza {{variable}} en párrafos y celdas de tabla.
    Reconstruye el texto completo de cada párrafo para manejar
    variables fragmentadas en múltiples runs.
    """
    def procesar(parrafo):
        texto = "".join(r.text for r in parrafo.runs)
        if "{{" not in texto:
            return
        for clave, valor in variables.items():
            texto = texto.replace(f"{{{{{clave}}}}}", str(valor or ""))
        if parrafo.runs:
            parrafo.runs[0].text = texto
            for run in parrafo.runs[1:]:
                run.text = ""

    for p in doc.paragraphs:
        procesar(p)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    procesar(p)
    return doc


def _docx_a_pdf(doc: Document) -> bytes:
    """
    Convierte un Document a PDF.
    Prioridad: LibreOffice (mantiene formato original) → reportlab (fallback de texto plano).
    """
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    try:
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", os.path.dirname(tmp_path), tmp_path],
            capture_output=True, timeout=30
        )
        pdf_path = tmp_path.replace(".docx", ".pdf")
        if os.path.exists(pdf_path):
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()
            os.unlink(pdf_path)
            return pdf_bytes
    except Exception:
        pass
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

    # Fallback — texto plano con reportlab
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 80
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "UNIVERSIDAD TECNOLÓGICA PASO DEL NORTE")
    y -= 25
    c.setFont("Helvetica", 10)
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        texto = p.text.strip()
        while len(texto) > 90:
            c.drawString(50, y, texto[:90])
            texto = texto[90:]
            y -= 14
            if y < 60:
                c.showPage(); y = height - 60
        c.drawString(50, y, texto)
        y -= 14
        if y < 60:
            c.showPage(); y = height - 60
    c.save()
    return buffer.getvalue()


def _aplicar_sello(pdf_bytes: bytes, x: int = 350, y: int = 100) -> bytes:
    """Superpone el sello PNG sobre la primera página del PDF."""
    sello_path = _ruta("sello.png")
    if not os.path.exists(sello_path):
        return pdf_bytes

    sello_buf = io.BytesIO()
    c = canvas.Canvas(sello_buf, pagesize=letter)
    c.drawImage(sello_path, x, y, width=120, height=120, mask="auto", preserveAspectRatio=True)
    c.save()
    sello_buf.seek(0)

    reader_doc   = PdfReader(io.BytesIO(pdf_bytes))
    reader_sello = PdfReader(sello_buf)
    writer       = PdfWriter()
    sello_page   = reader_sello.pages[0]

    for i, page in enumerate(reader_doc.pages):
        if i == 0:
            page.merge_page(sello_page)
        writer.add_page(page)

    resultado = io.BytesIO()
    writer.write(resultado)
    return resultado.getvalue()


def _nombre_alumno(alumno: dict) -> str:
    return f"{alumno['first_name']} {alumno['paternal_surname']} {alumno.get('maternal_surname', '')}".strip()


def generar_constancia(alumno: dict, con_sello: bool = False) -> bytes:
    """Constancia de servicio social — 1er cuatrimestre o general."""
    template = "constancia_1er_cuatri.docx" if alumno.get("cuatrimestre") == 1 else "constancia_general.docx"
    doc = _reemplazar_variables(Document(_ruta(template)), {
        "nombre_alumno":      _nombre_alumno(alumno),
        "matricula":          alumno["student_id"],
        "carrera":            alumno.get("career", ""),
        "programa":           alumno.get("actividad", ""),
        "periodo_inicio":     alumno.get("periodo_inicio", "Enero 2026"),
        "periodo_fin":        alumno.get("periodo_fin", "Abril 2026"),
        "fecha":              datetime.now().strftime("%d de %B de %Y"),
        "sello_departamento": "",
        "sello_institucion":  "",
    })
    pdf = _docx_a_pdf(doc)
    return _aplicar_sello(pdf) if con_sello else pdf


def generar_boleta(alumno: dict, con_sello: bool = False) -> bytes:
    """Boleta de liberación del servicio social."""
    doc = _reemplazar_variables(Document(_ruta("boleta_liberacion.docx")), {
        "nombre_alumno":      _nombre_alumno(alumno),
        "matricula":          alumno["student_id"],
        "grupo":              alumno.get("grupo", ""),
        "programa":           alumno.get("actividad", ""),
        "periodo_inicio":     alumno.get("periodo_inicio", "Enero 2026"),
        "periodo_fin":        alumno.get("periodo_fin", "Abril 2026"),
        "sello_temporal":     "",
        "sello_departamento": "",
    })
    pdf = _docx_a_pdf(doc)
    return _aplicar_sello(pdf) if con_sello else pdf


def generar_reporte_final(alumno: dict, form: dict) -> bytes:
    """Reporte final — lo llena el alumno."""
    doc = _reemplazar_variables(Document(_ruta("reporte_final.docx")), {
        "nombre_alumno":  _nombre_alumno(alumno),
        "matricula":      alumno["student_id"],
        "grupo":          alumno.get("grupo", ""),
        "telefono":       form.get("telefono", ""),
        "institucion":    form.get("institucion", ""),
        "fecha_inicio":   form.get("fecha_inicio", ""),
        "fecha_fin":      form.get("fecha_fin", ""),
        "actividades":    form.get("actividades", ""),
        "sello_temporal": "",
    })
    return _docx_a_pdf(doc)


def generar_control_horas(alumno: dict, form: dict) -> bytes:
    """Control de horas — lo llena el alumno."""
    doc = _reemplazar_variables(Document(_ruta("control_horas.docx")), {
        "nombre_alumno":   _nombre_alumno(alumno),
        "matricula":       alumno["student_id"],
        "telefono":        form.get("telefono", ""),
        "institucion":     form.get("institucion", ""),
        "fecha_inicio":    form.get("fecha_inicio", ""),
        "fecha_fin":       form.get("fecha_fin", ""),
        "programa":        alumno.get("actividad", form.get("programa", "")),
        "hora_entrada":    form.get("hora_entrada", ""),
        "hora_salida":     form.get("hora_salida", ""),
        "horas_acumuladas": str(alumno.get("accumulated_hours", "")),
        "fecha_asistencia": "",
        "horas_dia":        "",
        "sello_temporal":   "",
    })
    return _docx_a_pdf(doc)


def generar_evaluacion_desempeno(alumno: dict, form: dict) -> bytes:
    """Evaluación de desempeño — la llena el encargado."""
    def x(campo, opcion): return "X" if form.get(campo) == opcion else ""

    doc = _reemplazar_variables(Document(_ruta("evaluacion_desempeno.docx")), {
        "nombre_alumno":               _nombre_alumno(alumno),
        "matricula":                   alumno["student_id"],
        "carrera":                     alumno.get("career", ""),
        "fecha":                       form.get("fecha", datetime.now().strftime("%d/%m/%Y")),
        "institucion":                 form.get("institucion", ""),
        "direccion":                   form.get("direccion", ""),
        "tel_institucion":             form.get("tel_institucion", ""),
        "supervisor":                  form.get("supervisor", ""),
        "días_modalidad":              form.get("dias_modalidad", ""),
        "cumple_horario_siempre":      x("cumple_horario", "siempre"),
        "cumple_horario_algunas_veces": x("cumple_horario", "algunas_veces"),
        "cumple_horario_nunca":        x("cumple_horario", "nunca"),
        "supervisa_entrada_salida_si": x("supervisa_entrada_salida", "si"),
        "supervisa_entrada_salida_no": x("supervisa_entrada_salida", "no"),
        "firma_horas_semanales_si":    x("firma_horas_semanales", "si"),
        "firma_horas_semanales_no":    x("firma_horas_semanales", "no"),
        "reporte_final_si":            x("reporte_final", "si"),
        "reporte_final_no":            x("reporte_final", "no"),
        "acato_reglamento_si":         x("acato_reglamento", "si"),
        "acato_reglamento_no":         x("acato_reglamento", "no"),
        "acato_reglamento_porque":     form.get("acato_reglamento_porque", ""),
        "labores_eficientes_si":       x("labores_eficientes", "si"),
        "labores_eficientes_no":       x("labores_eficientes", "no"),
        "labores_eficientes_porque":   form.get("labores_eficientes_porque", ""),
        "actividades_relevantes":      form.get("actividades_relevantes", ""),
        "nivel_relevancia_relevantes": x("nivel_relevancia", "relevantes"),
        "nivel_relevancia_poco":       x("nivel_relevancia", "poco_relevantes"),
        "nivel_relevancia_nada":       x("nivel_relevancia", "nada_relevantes"),
        "retribuye_formacion_mucho":   x("retribuye_formacion", "mucho"),
        "retribuye_formacion_poco":    x("retribuye_formacion", "poco"),
        "retribuye_formacion_nada":    x("retribuye_formacion", "nada"),
        "calificación_1":              x("calificacion", "1"),
        "calificación_2":              x("calificacion", "2"),
        "calificación_3":              x("calificacion", "3"),
        "calificación_4":              x("calificacion", "4"),
        "calificación_5":              x("calificacion", "5"),
        "calificacion_porque":         form.get("calificacion_porque", ""),
        "sello_institucion":           "",
    })
    return _docx_a_pdf(doc)
